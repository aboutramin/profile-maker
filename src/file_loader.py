from pathlib import Path

from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".tex",
    ".json",
    ".csv",
    ".docx",
}


def load_text_file(path: str) -> str:
    file_path = Path(path)
    if not file_path.exists():
        return ""
    return file_path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: str) -> str:
    file_path = Path(path)
    if not file_path.exists():
        return ""

    reader = PdfReader(str(file_path))
    pages = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"\n\n--- PAGE {index} ---\n{text}")

    return "\n".join(pages)


def load_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to read .docx files. "
            "Install dependencies with: pip install -r requirements.txt"
        ) from exc

    file_path = Path(path)
    if not file_path.exists():
        return ""

    document = Document(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))

    return "\n".join(paragraphs + table_rows)


def is_supported_document(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS


def load_document(path: str) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(path)

    if suffix == ".docx":
        return load_docx(path)

    return load_text_file(path)


def save_text(path: str, content: str) -> None:
    file_path = Path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_text(content, encoding="utf-8")
